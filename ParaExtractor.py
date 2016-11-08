import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pprint import pprint
import textwrap
# from pathlib import Path
# import pandas as pd

class ParaExtractor():
    def __init__(self, para_file, testing=False):
        self.doc = docx.Document(str(para_file))
        self.testing = testing

        if self.testing:
            print(len(self.doc.paragraphs), " docx paragraphs loaded")
    
            for par in self.doc.paragraphs:
                print(par.text)

        print('{paragraphs} docx paragraphs loaded'.format(paragraphs=len(self.doc.paragraphs)))
        self.paras = dict(self.get_paras())
        print('{paras} paras found'.format(paras=len(self.paras)))

        
    def relevant_paragraph(self, par):
        if par.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return False
        if len(par.text.strip()) == 0:
            return False
        return True

    def test_relevance(self):
        for par in self.doc.paragraphs:
            if self.relevant_paragraph(par):
                print("------PARAGRAPH-----\n", par.text)
            else:
                print("=====JUNK====\n", par.text)

    def extract_paras(self):
        para_list = self.doc.paragraphs
        para_start = re.compile("^\s*\((?P<page>\d+)-(?P<instance>\d+)\) (?P<text>.+)$")
        para_continue = re.compile("^\s*\(Continued from (?P<continuance>[^)]+)\) (?P<text>.+)$")
        current_para = None
        current_components = []
        for par in para_list:
            if self.relevant_paragraph(par):
                start_match = para_start.match(par.text)
                continue_match = para_continue.match(par.text)
                if start_match:
                    current_para = (int(start_match.group('page')), int(start_match.group('instance')))
                    para_text = start_match.group('text')
                    if self.testing:
                        print("\nSTARTING para {cp}:".format(cp=current_para))
                elif continue_match:
                    para_text = continue_match.group('text')
                    if self.testing:
                        print("CONTINUING {cp} from {c}:".format(cp=current_para, c=continue_match.group('continuance')))
                else:
                    para_text = par.text.strip()
                
                if self.testing:
                    print("> ", para_text)
                yield (current_para, para_text)
            else:
                if self.testing:
                    print('...[{x}]...'.format(x=par.text))
                  
    def text_extraction(self):
        pprint(list(self.extract_paras()))  

    def get_paras(self):
        current_para_id = None
        current_paragraphs = None
        for para_id, paragraph in self.extract_paras():
            if current_para_id != para_id:  # Moving to a new element
                if current_para_id:
                    yield (current_para_id, '\n'.join(current_paragraphs))
                current_para_id = para_id
                current_paragraphs = [paragraph]
            else:
                if current_para_id:
                    current_paragraphs.append(paragraph)
                else:
                    print('Orphaned text dropped: ', paragraph)
        # And spit out the leftovers
        yield (current_para_id, '\n'.join(current_paragraphs))

    def test_get_paras(self):
        parsed_paras = dict(self.get_paras())

        pprint(parsed_paras)

        for pid, text in sorted(parsed_paras.items()):
            print('------{p}------'.format(p=pid))
            print(textwrap.indent(text, '> '))